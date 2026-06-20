from docx import Document
import pandas as pd
import os

# Load dataset
df = pd.read_csv("data/customer_analytics.csv")

# Create document
document = Document()
document.add_heading("DATASET CHARACTERISTICS REPORT", level=1)

document.add_paragraph(f"Total Rows: {df.shape[0]}")
document.add_paragraph(f"Total Columns: {df.shape[1]}")

document.add_heading("Column Types & Missing Values", level=2)

null_counts = df.isnull().sum()

for col in df.columns:
    missing = null_counts[col]
    percentage = (missing / len(df)) * 100
    document.add_paragraph(
        f"{col} | {df[col].dtype} | Missing: {missing} ({percentage:.1f}%)"
    )

document.add_heading("Basic Statistics (Numerical Columns)", level=2)

stats = df.describe().transpose()
document.add_paragraph(stats.to_string())

document.add_heading("Duplicate Rows Check", level=2)

duplicates = df.duplicated().sum()
document.add_paragraph(f"Number of duplicate rows found: {duplicates}")

# Save file
document.save("data_summary_report.docx")

print("DOCX Report Generated Successfully!")