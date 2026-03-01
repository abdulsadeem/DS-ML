# ================================
# DATA SUMMARY REPORT GENERATOR
# ================================

import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# -------------------------------
# 1️⃣ Set Paths
# -------------------------------
dataset_path = r"C:\AIml_intern\data\customer_analytics.csv"   # 🔁 Change if needed
report_folder = r"C:\AIml_intern\reports"

# Create reports folder if not exists
os.makedirs(report_folder, exist_ok=True)

# -------------------------------
# 2️⃣ Load Dataset
# -------------------------------
df = pd.read_csv(dataset_path)

# ===============================
# Phase 1 – Data Inspection
# ===============================

from io import StringIO

# Capture df.info()
buffer = StringIO()
df.info(buf=buffer)
info_text = buffer.getvalue()

# First 5 rows
head_text = df.head().to_string()

# Statistical summary
describe_text = df.describe().to_string()
# -------------------------------
# 3️⃣ Phase 2 – Data Cleaning Checks
# -------------------------------
missing_values = df.isnull().sum()
duplicate_rows = df[df.duplicated()]
duplicate_count = len(duplicate_rows)

# -------------------------------
# 4️⃣ Phase 3 – Generate Plots
# -------------------------------

# ---- Age Distribution ----
plt.figure(figsize=(8,5))
df['Age'].hist(bins=15, color='skyblue', edgecolor='black')
plt.title("Age Distribution")
plt.xlabel("Age")
plt.ylabel("Number of Customers")
plt.tight_layout()
plt.savefig(os.path.join(report_folder, "age_distribution.png"))
plt.close()

# ---- Income Distribution ----
plt.figure(figsize=(8,5))
sns.histplot(df['AnnualIncome'], kde=True, color='green', bins=15)
plt.title("Annual Income Distribution")
plt.xlabel("Annual Income")
plt.ylabel("Number of Customers")
plt.tight_layout()
plt.savefig(os.path.join(report_folder, "income_distribution.png"))
plt.close()

# ---- Gender Distribution ----
plt.figure(figsize=(6,4))
sns.countplot(x='Gender', data=df, color='skyblue')
plt.title("Gender Distribution")
plt.tight_layout()
plt.savefig(os.path.join(report_folder, "gender_distribution.png"))
plt.close()

# ---- Income vs Spending ----
plt.figure(figsize=(8,5))
sns.scatterplot(x='AnnualIncome', y='SpendingScore', data=df, hue='Gender')
plt.title("Income vs Spending Score")
plt.xlabel("Annual Income")
plt.ylabel("Spending Score")
plt.tight_layout()
plt.savefig(os.path.join(report_folder, "income_vs_spending.png"))
plt.close()

# ---- Spending by Gender ----
plt.figure(figsize=(8,5))
sns.boxplot(x='Gender', y='SpendingScore', data=df, color='lightgreen')
# ===============================
# Phase 4 – Top Insights
# ===============================

correlation_matrix = df.corr(numeric_only=True)

# Get top correlations (excluding self-correlation)
corr_pairs = (
    correlation_matrix.unstack()
    .sort_values(ascending=False)
)

# Remove duplicate and self correlations
corr_pairs = corr_pairs[corr_pairs < 1]

top_3 = corr_pairs.head(3)

insights_text = ""
for (var1, var2), value in top_3.items():
    insights_text += f"{var1} and {var2} have correlation of {round(value, 2)}\n"


# ---- Correlation Heatmap ----
plt.figure(figsize=(10,6))
sns.heatmap(df.corr(numeric_only=True), annot=True, cmap="coolwarm")
plt.title("Correlation Heatmap")
plt.tight_layout()
plt.savefig(os.path.join(report_folder, "correlation_heatmap.png"))
plt.close()

# -------------------------------
# 5️⃣ Generate DOCX Report
# -------------------------------

doc = Document()
doc.add_heading("Data Summary Report", 0)

# ===============================
# Phase 1 – Documentation
# ===============================

doc.add_heading("Phase 1: Data Inspection", level=1)

doc.add_heading("Dataset Preview (First 5 Rows)", level=2)
doc.add_paragraph(head_text)

doc.add_heading("Dataset Structure (df.info())", level=2)
doc.add_paragraph(info_text)

doc.add_heading("Statistical Summary (df.describe())", level=2)
doc.add_paragraph(describe_text)

# Phase 2
doc.add_heading("Phase 2: Data Cleaning", level=1)

doc.add_heading("Missing Values", level=2)
for col, count in missing_values.items():
    doc.add_paragraph(f"{col}: {count}")

doc.add_heading("Duplicate Rows", level=2)
doc.add_paragraph(f"Total duplicate rows found: {duplicate_count}")

# Phase 3 – Univariate
doc.add_heading("Phase 3: Univariate Analysis", level=1)

doc.add_heading("Age Distribution", level=2)
doc.add_paragraph("Most customers fall between age 25 and 40.")
doc.add_picture(os.path.join(report_folder, "age_distribution.png"), width=Inches(5))

doc.add_heading("Annual Income Distribution", level=2)
doc.add_paragraph("Income distribution shows slight right skew.")
doc.add_picture(os.path.join(report_folder, "income_distribution.png"), width=Inches(5))

doc.add_heading("Gender Distribution", level=2)
doc.add_picture(os.path.join(report_folder, "gender_distribution.png"), width=Inches(5))

# Bivariate
doc.add_heading("Phase 3: Bivariate Analysis", level=1)

doc.add_heading("Income vs Spending Score", level=2)
doc.add_paragraph("Higher income customers tend to have higher spending scores.")
doc.add_picture(os.path.join(report_folder, "income_vs_spending.png"), width=Inches(5))

doc.add_heading("Spending Score by Gender", level=2)
doc.add_paragraph("Spending patterns vary slightly across gender groups.")
doc.add_picture(os.path.join(report_folder, "spending_by_gender.png"), width=Inches(5))

# Phase 4
doc.add_heading("Phase 4: Multivariate Analysis", level=1)

doc.add_heading("Correlation Heatmap", level=2)
doc.add_paragraph("The heatmap shows relationships between numerical variables.")
doc.add_picture(os.path.join(report_folder, "correlation_heatmap.png"), width=Inches(5))

# Save Report
doc_path = os.path.join(report_folder, "data_summary_report.docx")
doc.save(doc_path)

print("✅ Report generated successfully!")
print(f"📄 Saved at: {doc_path}")